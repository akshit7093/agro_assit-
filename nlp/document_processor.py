import os
import logging
from typing import Dict, Any, List, Optional
import docx

logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text content from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        str: Extracted text content
    """
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text:
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        full_text.append(cell.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX file {file_path}: {str(e)}")
        return f"Error processing document: {str(e)}"

def process_document(file_path: str) -> Dict[str, Any]:
    """
    Process a document file and extract its content.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Dict containing the extracted content and metadata
    """
    if not os.path.exists(file_path):
        return {"error": f"File not found: {file_path}"}
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_ext == '.docx':
            content = extract_text_from_docx(file_path)
            return {
                "content": content,
                "file_path": file_path,
                "file_type": "docx",
                "file_name": os.path.basename(file_path)
            }
        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return {
                "content": content,
                "file_path": file_path,
                "file_type": "txt",
                "file_name": os.path.basename(file_path)
            }
        elif file_ext in ['.pdf']:
            # Placeholder for PDF processing
            return {"error": "PDF processing not implemented yet"}
        else:
            return {"error": f"Unsupported file type: {file_ext}"}
    except Exception as e:
        logger.error(f"Error processing document {file_path}: {str(e)}")
        return {"error": f"Error processing document: {str(e)}"}